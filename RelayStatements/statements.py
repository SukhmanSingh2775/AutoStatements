from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium import webdriver
from docx.enum.text import WD_ALIGN_PARAGRAPH 
import time

import xlrd
import openpyxl


from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt





browser = webdriver.Chrome("/Users/sukhmansingh/Desktop/Python/chromedriver")
browser.set_window_size(1920, 1080)
browser.get("https://relay.amazon.com/")
time.sleep(2)

signInBtn = browser.find_element_by_link_text('Sign In')
signInBtn.click()

gold_star_username = 'ranjitsingh9774@gmail.com'
akal_trucking_username = 'harman4985@gmail.com'

gold_star_password = "Singh20"
akal_trucking_password = "Amazon4980"

account = int(input("Which Account? Goldstar (1) , Akal (2)"))

if account == 1:
    username = gold_star_username
    password = gold_star_password
    
else:
    username = akal_trucking_username
    password = akal_trucking_password


emailField = browser.find_element_by_id('ap_email')
emailField.click()
emailField.send_keys(username)
time.sleep(0.5)

passwordField = browser.find_element_by_id('ap_password')
passwordField.click()
passwordField.send_keys(password)


time.sleep(1)


browser.find_element_by_id('signInSubmit').click()
number_of_statements = int(input("How many times do you to run it?"))



time.sleep(2)

browser.get('https://relay.amazon.com/tours/tours?state=history')

time.sleep(1)

    
final_load_id = []

def getStatements():
    
    document = Document()
    
    ready = input("Start?")
        
    driverSelection = int(input("Select Driver Ranjit Singh (1) , Harmanjit Singh (2), Sarabjit Singh (3), New Driver (4)"))
    driverName = ""
    search_bar = browser.find_element_by_xpath('/html/body/div[22]/main/div[5]/div/div/div[1]/div/div[1]/input')
    
    
    if driverSelection == 1:
        driverName = "Ranjit Singh"
        search_bar.click()
        search_bar.send_keys(driverName)
        
        
    elif driverSelection == 2:
        driverName = "Harmanjit Singh"
        search_bar.click()    
        search_bar.send_keys(driverName)
        
    elif driverSelection == 3:
        driverName = "Sarabjit Singh"
        search_bar.click()
        search_bar.send_keys(driverName)
        
    elif driverSelection == 4:
        driverName = input("What is the new driver name?")
        search_bar.click()
        search_bar.send_keys(driverName)
        
        
    time.sleep(4)
    
    cities = []
    prices = []
    loadId = []
    
    cities.clear()
    prices.clear()
    loadId.clear()
    
    final_prices = []
    final_cities = []
    final_load_id = []
    
    final_prices.clear()
    final_cities.clear()
    final_load_id.clear()

    cities = browser.find_elements_by_class_name("city")
    prices = browser.find_elements_by_class_name('tour-header__payout--normal--primary')
    loadId = browser.find_elements_by_class_name('tour-header__tour-id__sliced')
    
    times = browser.find_elements_by_class_name('tour-header__secondary ')

    dateRange = browser.find_element_by_xpath('//*[@id="application"]/div/div/span/div/div/span').text


    final_cities = []
    final_prices = []
    final_load_id = []
    final_times = []
 

    driverHeading = document.add_heading('',0)
    run = driverHeading.add_run("Driver Name: {}".format(driverName))
    font = run.font 
    font.color.rgb = RGBColor(255, 132, 0)
    driverHeading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run.italic = True


    dateRangeHeading = document.add_heading('',2)
    run = dateRangeHeading.add_run(dateRange)
    font = run.font
    font.color.rgb = RGBColor(170, 0, 255)

    dateRangeHeading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
   
    
    companyName = ''
    
    if (username == gold_star_username):
        companyName = 'Goldstar Trucking'
        companyNameHeading = document.add_heading('',2)
        run = companyNameHeading.add_run(companyName)
        font = run.font
        font.color.rgb = RGBColor(170, 0, 255)
    
    
    else:
        companyName = 'Akal Trucking'
        companyNameHeading = document.add_heading('',2)
        run = companyNameHeading.add_run(companyName)
        font = run.font
        font.color.rgb = RGBColor(170, 0, 255)


    document.add_heading('',2)

    for i in range(0, len(cities)):
        try:
            final_cities.append(cities[i].text)
            
            final_prices.append(prices[i].text)
            
            final_load_id.append(loadId[i].text)

        
            
            
        except:
            pass
            
    
    final_prices = [s.replace(',','') for s in final_prices]
    final_prices = [s for s in final_prices if s]

    final_load_id = [s.replace('.','') for s in final_load_id]
    final_load_id = [s[1:] if len(s) > 4 else s for s in final_load_id]
    
    try:
        final_times = [times[i].text for i in range(len(times))]
        final_times = [i[4:] for i in final_times if i]
        final_times = [i[:-4] for i in final_times if i]
    
        final_times_list = [i for i in final_times if len(i) >= 10]

    except:
        pass

    



    x = 0
    a = 0
    
    final_cities = [s for s in final_cities if s]
    for i in range(len(final_cities) // 2):

            load_paragraph = document.add_paragraph(style="List Number")
        
            try:
                
                run = load_paragraph.add_run("({})".format(final_times_list[x]) + "   " + "({})".format(final_times_list[x+1]) + " ")
                font = run.font
                font.color.rgb = RGBColor(170, 0, 255)
                run.bold = True

                
            

                run = load_paragraph.add_run(final_cities[x] + "   ")
                font = run.font
                font.color.rgb = RGBColor(33, 83, 163)
                run.bold = True
                
                run = load_paragraph.add_run("To" + "   ")
                font = run.font
                font.color.rgb = RGBColor(10,20,30)
                run.bold = True
                
                
                run = load_paragraph.add_run(final_cities[x+1] + "   ")
                font = run.font
                font.color.rgb = RGBColor(33, 83, 163)
                run.bold = True
                
                
                
                run = load_paragraph.add_run(final_prices[i] + "   ")
                font = run.font
                font.color.rgb = RGBColor(247, 21, 5)
                run.bold = True
                
                
                
                
                run = load_paragraph.add_run("#" + final_load_id[i])
                font = run.font
                font.color.rgb = RGBColor(247, 5, 239)
                run.bold = True
                
            
                
                
                x += 2
            
            except:
                pass
        
        
        
            document.add_paragraph('')



    lst = [e[1:] for e in final_prices]

    lst = [float(i) for i in lst]

    total = sum(lst)


    total_paragraph = document.add_paragraph()
    run = total_paragraph.add_run("Total:" + "  " + "$" + str(int(total)))
    font = run.font
    run.bold = True
    run.italic = True
    run.underline = True
    font.color.rgb = RGBColor(0, 255, 183)
    total_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font.size = Pt(24)
    
    file_path_string = '{}.docx'.format(driverName +  ' ' + "{}".format(dateRange))
    path = '/Users/sukhmansingh/Desktop/Invoices2/'
    
    if (username == gold_star_username):
        path = path + "Goldstar/{}/".format(driverName)
        document.save(path  + file_path_string)
  
    
    else:
        path = path + "Akal/{}/".format(driverName) 
        document.save(path  + file_path_string)
    


for i in range(0, number_of_statements):
    getStatements()
    
    
    
    
    
    
    
    
    
    
# path = ("/Users/sukhmansingh/Downloads/ExcelStatement1.xlsx")

# wb = xlrd.open_workbook(path)
# sheet = wb.sheet_by_index(1)

# load_list_1 = sheet.col_values(1) #Unmodified round trip loads
# load_list_2 = sheet.col_values(2) #Unmodified one way laods
# prices_column = sheet.col_values(21) # Unmodified prices

# round_trips = [i[-4:] for i in load_list_1  if i != "Trip ID"] # Round Trips Loads
# load_list_2 = [i[-4:] for i in load_list_2 if i != "Load ID"] # One Way Loads
# prices_column = [i for i in prices_column if type(i) != str] # Getting float only for prices
# prices_column.pop(-1)

# one_way = [load_list_2[i] for i in range(len(round_trips)) if round_trips[i] == "" or round_trips[i] == "-" if load_list_2[i] != ""] #Modified one way loads

# prices_one_way = [prices_column[i] for i in range(len(prices_column)) if load_list_1[i] == "" or load_list_1[i] == "-"]
# prices_round_trip = [prices_column[i] for i in range(len(prices_column)) if load_list_1[i] and load_list_1[i] != "-"]
# print("Round trips are " + str(round_trips))
# print("One way loads are " + str(one_way))
# print("Final load id is " + str(final_load_id))
            
# for i in range(len(final_load_id)):
#     if final_load_id[i] in one_way or final_load_id[i] in round_trips:
#         if (final_load_id[i] in one_way):
#             print("Found it in one way")
#         elif (final_load_id[i] in round_trips):
#             print("Found it in round trips")










    
    
    
    



# def getCancelledLoads():
    
#     document = Document()
#     search_bar = browser.find_element_by_xpath('/html/body/div[22]/main/div[5]/div/div/div[1]/div/div[1]/input')
#     search_bar.click()
#     time.sleep(1)
#     clear = browser.find_element_by_xpath('//*[@id="application"]/div/div/div[1]/div/div[1]/i')
#     clear.click()
#     time.sleep(2)
    
#     filter_btn = browser.find_element_by_xpath('//*[@id="application"]/div/div/div[1]/div/div[2]/div[2]/span/span')
    
#     filter_btn.click()
    
#     cancelled_btn = browser.find_element_by_xpath('//*[@id="application"]/div/div/div[1]/div/div[2]/div[2]/div/div[9]/span')
#     cancelled_btn.click()
    
    
#     time.sleep(2)
    
#     cities = browser.find_elements_by_class_name("city")
    
#     prices = browser.find_elements_by_class_name('tour-header__payout--normal--primary')
#     loadId = browser.find_elements_by_class_name('tour-header__tour-id__sliced')
#     times = browser.find_elements_by_class_name('tour-header__secondary ')
#     dateRange = browser.find_element_by_xpath('//*[@id="application"]/div/div/span/div/div/span').text


#     final_cities = []
#     final_prices = []
#     final_load_id = []
#     final_times = []


#     dateRangeHeading = document.add_heading('',2)
#     run = dateRangeHeading.add_run(dateRange)
#     font = run.font
#     font.color.rgb = RGBColor(170, 0, 255)

#     dateRangeHeading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
   
    
#     companyName = ''
    
#     if (username == gold_star_username):
#         companyName = 'Goldstar Trucking'
#         companyNameHeading = document.add_heading('',2)S
#         run = companyNameHeading.add_run(companyName)
#         font = run.font
#         font.color.rgb = RGBColor(170, 0, 255)
    
    
#     else:
#         companyName = 'Akal Trucking'
#         companyNameHeading = document.add_heading('',2)
#         run = companyNameHeading.add_run(companyName)
#         font = run.font
#         font.color.rgb = RGBColor(170, 0, 255)


#     document.add_heading('',2)

#     for i in range(0, len(cities)):
#         try:
#             final_cities.append(cities[i].text)
        
#             final_prices.append(prices[i].text)
            
#             final_load_id.append(loadId[i].text)

        
            
            
#         except:
#             pass
            



    
#     final_prices = [s.replace(',','') for s in final_prices]

#     final_load_id = [s.replace('.','') for s in final_load_id]
    
#     final_load_id = [s[1:] if len(s) > 4 else s for s in final_load_id]
    
        
#     try:
#         final_times = [times[i].text for i in range(len(times))]
#         final_times = [i[4:] for i in final_times if i]
#         final_times = [i[:-4] for i in final_times if i]
    
#         final_times_list = [i for i in final_times if len(i) >= 10]

#     except:
#         pass



#     x = 0
    
    

#     for i in range(len(final_cities) // 2):

#             load_paragraph = document.add_paragraph(style="List Number")
        
#             try:
                
#                 run = load_paragraph.add_run("({})".format(final_times_list[x]) + "   " + "({})".format(final_times_list[x+1]) + " ")
#                 font = run.font
#                 font.color.rgb = RGBColor(170, 0, 255)
#                 run.bold = True

                
                

#                 run = load_paragraph.add_run(final_cities[x] + "   ")
#                 font = run.font
#                 font.color.rgb = RGBColor(33, 83, 163)
#                 run.bold = True
                
#                 run = load_paragraph.add_run("To" + "   ")
#                 font = run.font
#                 font.color.rgb = RGBColor(10,20,30)
#                 run.bold = True
                
                
#                 run = load_paragraph.add_run(final_cities[x+1] + "   ")
#                 font = run.font
#                 font.color.rgb = RGBColor(33, 83, 163)
#                 run.bold = True
                
                
                
#                 run = load_paragraph.add_run(final_prices[i] + "   ")
#                 font = run.font
#                 font.color.rgb = RGBColor(247, 21, 5)
#                 run.bold = True
                
                
                
                
#                 run = load_paragraph.add_run("#" + final_load_id[i])
#                 font = run.font
#                 font.color.rgb = RGBColor(247, 5, 239)
#                 run.bold = True
                
            
                
                
#                 x += 2
            
#             except:
#                 pass
        
        
        
#             document.add_paragraph('')





#     lst = [e[1:] for e in final_prices]

#     lst = [float(i) for i in lst]

#     total = sum(lst)


#     total_paragraph = document.add_paragraph()
#     run = total_paragraph.add_run("Total:" + "  " + "$" + str(int(total)))
#     font = run.font
#     run.bold = True
#     run.italic = True
#     run.underline = True
#     font.color.rgb = RGBColor(0, 255, 183)
#     total_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
#     font.size = Pt(24)

    
    
    
    
    
    
    
    
#     path = '/Users/sukhmansingh/Desktop/Invoices2/'
    
    
#     if (username == gold_star_username):
#         path = path + "Goldstar/"
#         document.save(path + '{}.docx'.format("CancelledLoads" +  ' ' + "{}".format(dateRange)))
  
    
#     else:
#         path = path + "Akal/"
#         document.save(path + '{}.docx'.format("CancelledLoads" +  ' ' + "{}".format(dateRange)))


# run_cancel = input("Do you want to get the cancel statement? ")

# if run_cancel:
#     getCancelledLoads()
